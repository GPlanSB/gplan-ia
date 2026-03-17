import streamlit as st
import google.generativeai as genai
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document

# Configuração de Estilo "Elon Musk" (Minimalista e Dark)
st.set_page_config(page_title="GPlan 4.0", page_icon="🎯", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #0b0e11; }
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; background-color: #1E1E1E; color: white; border: 1px solid #333; }
    .stTextInput>div>div>input { background-color: #121212; color: white; border: 1px solid #333; }
    </style>
    """, unsafe_allow_html=True)

# Configuração da API
# Substitua pelo seu segredo no Streamlit ou variável de ambiente
GEMINI_API_KEY = st.sidebar.text_input("Gemini API Key", type="password")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-pro')

# --- FUNÇÕES CORE ---
def gerar_relatorio_docx(conteudo):
    doc = Document()
    doc.add_heading('GPlan 4.0 - Relatório Executivo', 0)
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(conteudo)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- INTERFACE ---
st.title("🎯 GPlan 4.0")
st.caption("Agente Especialista em Gerenciamento de Projetos Master | High-Performance Strategy")

with st.sidebar:
    st.header("⚙️ Central de Controle")
    uploaded_file = st.file_uploader("Upload de Dados (Excel/CSV)", type=['xlsx', 'csv'])
    modo = st.selectbox("Modo de Operação", ["Consultoria Estratégica", "Geração de Cronograma", "Análise de Riscos", "Relatório de Status"])
    st.divider()
    st.info("O GPlan 4.0 utiliza raciocínio avançado para prever gargalos antes que eles ocorram.")

# Inicializar Chat
if "messages" not in st.session_state:
    st.session_state.messages = []

# Exibir mensagens
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# --- LÓGICA DE AGENTE ---
if prompt := st.chat_input("Comande o GPlan 4.0..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        contexto_dados = ""
        if uploaded_file:
            df = pd.read_excel(uploaded_file) if uploaded_file.name.endswith('xlsx') else pd.read_csv(uploaded_file)
            contexto_dados = f"\nDados Atuais do Projeto:\n{df.to_string()}"

        # System Prompt de Especialista Master
        full_prompt = f"""
        Você é o GPlan 4.0, o coordenador de projetos mais avançado do mundo. 
        Sua personalidade: Consultor sênior, direto, focado em eficiência máxima (estilo Elon Musk).
        Suas capacidades: Domínio total de PMBOK, Lean, Six Sigma e Gestão Ágil.
        
        CONTEXTO DO PROJETO:
        {contexto_dados}
        
        INSTRUÇÃO DO USUÁRIO:
        {prompt}
        
        Se solicitado um cronograma, formate em tabela. Se solicitado um relatório, seja técnico e profundo.
        Sempre identifique riscos críticos ocultos.
        """

        response = model.generate_content(full_prompt)
        text_response = response.text
        st.markdown(text_response)
        
        # Opção de Download de Relatório se a IA gerou algo longo
        if len(text_response) > 500:
            docx_file = gerar_relatorio_docx(text_response)
            st.download_button(label="📄 Baixar Relatório Oficial (.docx)", 
                               data=docx_file, 
                               file_name=f"Relatorio_GPlan_{datetime.now().strftime('%Y%m%d')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.session_state.messages.append({"role": "assistant", "content": text_response})

